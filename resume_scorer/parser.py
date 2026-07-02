"""Resume and job-description text extraction — no API calls."""

from __future__ import annotations

import io
import re
from typing import Optional

import fitz  # PyMuPDF
from docx import Document


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def _page_text_column_ordered(page: fitz.Page) -> str:
    """
    Reassemble page text with left-column blocks before right-column blocks.
    Helps multi-column resumes (e.g. Jake's template) keep sidebar headers like
    Skills near their content instead of interleaved with main-column text.
    """
    blocks = page.get_text("blocks")
    if not blocks:
        return page.get_text()

    typed: list[tuple[float, float, str]] = []
    for block in blocks:
        if len(block) < 5:
            continue
        text = str(block[4]).strip()
        if not text:
            continue
        typed.append((float(block[0]), float(block[1]), text))

    if not typed:
        return page.get_text()

    page_width = page.rect.width
    # Jake's-style sidebar is ~30–38% from the left; use adaptive split when possible
    x0_values = sorted(b[0] for b in typed)
    split_x = page_width * 0.38
    if len(x0_values) >= 4:
        mid = len(x0_values) // 2
        gap = x0_values[mid] - x0_values[mid - 1]
        if gap > page_width * 0.08:
            split_x = (x0_values[mid - 1] + x0_values[mid]) / 2.0

    left = sorted((y, t) for x, y, t in typed if x < split_x)
    right = sorted((y, t) for x, y, t in typed if x >= split_x)
    ordered = [t for _, t in left] + [t for _, t in right]
    return "\n".join(ordered)


def extract_text_from_pdf(data: bytes) -> str:
    doc = fitz.open(stream=data, filetype="pdf")
    parts: list[str] = []
    for page in doc:
        parts.append(_page_text_column_ordered(page))
    doc.close()
    return "\n".join(parts).strip()


def extract_text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    parts: list[str] = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts).strip()


def extract_text_from_txt(data: bytes) -> str:
    return _decode_text(data).strip()


def extract_text(data: bytes, filename: str) -> str:
    """Extract plain text from uploaded file bytes based on extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return extract_text_from_pdf(data)
    if lower.endswith(".docx"):
        return extract_text_from_docx(data)
    if lower.endswith(".txt"):
        return extract_text_from_txt(data)
    raise ValueError(f"Unsupported file type: {filename}")


def detect_multi_column_pdf(data: bytes) -> bool:
    """True when PDF blocks suggest a multi-column layout."""
    doc = fitz.open(stream=data, filetype="pdf")
    try:
        for page in doc:
            blocks = page.get_text("blocks")
            x0_values = sorted(
                float(b[0]) for b in blocks if len(b) >= 5 and str(b[4]).strip()
            )
            if len(x0_values) < 4:
                continue
            page_width = page.rect.width
            mid = len(x0_values) // 2
            gap = x0_values[mid] - x0_values[mid - 1]
            if gap > page_width * 0.08:
                return True
    finally:
        doc.close()
    return False


def docx_has_tables(data: bytes) -> bool:
    """True when DOCX contains table structures."""
    doc = Document(io.BytesIO(data))
    return len(doc.tables) > 0


def alphanumeric_ratio(text: str) -> float:
    if not text:
        return 0.0
    alnum = sum(1 for c in text if c.isalnum())
    return alnum / len(text)


def validate_extracted_text(text: str) -> Optional[str]:
    """Return a warning message if text looks unreadable, else None."""
    if len(text) < 50:
        return (
            "Extracted text is very short (< 50 characters). "
            "The file may be image-based or unreadable by text extraction."
        )
    if alphanumeric_ratio(text) < 0.4:
        return (
            "Extracted text contains mostly non-alphanumeric characters. "
            "The file may be a scanned/image PDF without OCR."
        )
    return None


def guess_name_from_pdf(data: bytes) -> Optional[str]:
    """
    Guess candidate name from largest font-size text near the top of page 1.
    Uses PyMuPDF span-level font data; returns None if unavailable or inconclusive.
    """
    doc = fitz.open(stream=data, filetype="pdf")
    candidates: list[dict[str, float | str]] = []
    try:
        page = doc[0]
        page_height = page.rect.height
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                y = line["bbox"][1]
                if y > page_height * 0.35:
                    continue
                for span in line.get("spans", []):
                    text = span.get("text", "").strip()
                    if not text:
                        continue
                    candidates.append({
                        "text": text,
                        "size": float(span.get("size", 0)),
                        "x0": float(span["bbox"][0]),
                        "y": float(y),
                    })
    finally:
        doc.close()

    if not candidates:
        return None

    max_size = max(c["size"] for c in candidates)
    top_size = [c for c in candidates if c["size"] >= max_size - 0.5]
    by_line: dict[float, list[dict[str, float | str]]] = {}
    for c in top_size:
        by_line.setdefault(round(c["y"], 1), []).append(c)

    best_y = max(by_line, key=lambda y: max(c["size"] for c in by_line[y]))
    line_spans = sorted(by_line[best_y], key=lambda c: c["x0"])
    return " ".join(str(c["text"]) for c in line_spans).strip() or None
