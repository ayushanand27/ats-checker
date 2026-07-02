"""Structured JSON → DOCX via python-docx — no LLM."""

from __future__ import annotations

from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

# Template style constants — different look per template, same code path
TEMPLATE_STYLES: dict[str, dict[str, Any]] = {
    "jacks_tech": {
        "name_size": 18,
        "heading_size": 11,
        "body_size": 10,
        "name_color": RGBColor(0x2C, 0x3E, 0x50),
        "heading_color": RGBColor(0x2C, 0x3E, 0x50),
        "section_order": ["summary", "skills", "experience", "projects", "education", "achievements"],
    },
    "classic_nontech": {
        "name_size": 16,
        "heading_size": 12,
        "body_size": 11,
        "name_color": RGBColor(0x1A, 0x1A, 0x1A),
        "heading_color": RGBColor(0x33, 0x33, 0x33),
        "section_order": ["summary", "experience", "education", "skills", "projects", "achievements"],
    },
}


def _add_heading(doc: Document, text: str, style: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(style["heading_size"])
    run.font.color.rgb = style["heading_color"]
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)


def _add_bullet(doc: Document, text: str, style: dict[str, Any]) -> None:
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.size = Pt(style["body_size"])


def render_docx(data: dict[str, Any], template_name: str = "jacks_tech") -> bytes:
    """Build a DOCX resume from structured JSON."""
    style = TEMPLATE_STYLES.get(template_name, TEMPLATE_STYLES["jacks_tech"])
    doc = Document()
    sections = doc.sections[0]
    sections.top_margin = Inches(0.6)
    sections.bottom_margin = Inches(0.6)
    sections.left_margin = Inches(0.75)
    sections.right_margin = Inches(0.75)

    name = data.get("name", "Resume")
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(name)
    name_run.bold = True
    name_run.font.size = Pt(style["name_size"])
    name_run.font.color.rgb = style["name_color"]

    contact = data.get("contact", {})
    contact_parts = [
        contact.get("location", ""),
        contact.get("phone", ""),
        contact.get("email", ""),
        contact.get("linkedin", ""),
        contact.get("github", ""),
        contact.get("leetcode", ""),
        contact.get("portfolio", ""),
    ]
    contact_line = " | ".join(p for p in contact_parts if p)
    if contact_line:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = cp.add_run(contact_line)
        cr.font.size = Pt(style["body_size"])

    section_renderers = {
        "summary": lambda: _render_summary(doc, data, style),
        "skills": lambda: _render_skills(doc, data, style),
        "experience": lambda: _render_experience(doc, data, style),
        "education": lambda: _render_education(doc, data, style),
        "projects": lambda: _render_projects(doc, data, style),
        "achievements": lambda: _render_achievements(doc, data, style),
    }

    for section_key in style["section_order"]:
        section_renderers[section_key]()

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_summary(doc: Document, data: dict[str, Any], style: dict[str, Any]) -> None:
    summary = data.get("summary", "")
    if not summary:
        return
    _add_heading(doc, "Summary", style)
    p = doc.add_paragraph(summary)
    for run in p.runs:
        run.font.size = Pt(style["body_size"])


def _render_skills(doc: Document, data: dict[str, Any], style: dict[str, Any]) -> None:
    categories = data.get("skill_categories") or {}
    skills = data.get("skills", [])
    if isinstance(categories, dict) and any(categories.values()):
        _add_heading(doc, "Technical Skills", style)
        for label, items in categories.items():
            if not items:
                continue
            p = doc.add_paragraph()
            run = p.add_run(f"{label.replace('_', ' ').title()}: ")
            run.bold = True
            run.font.size = Pt(style["body_size"])
            run2 = p.add_run(", ".join(items))
            run2.font.size = Pt(style["body_size"])
        return
    if not skills:
        return
    _add_heading(doc, "Skills", style)
    p = doc.add_paragraph(", ".join(skills))
    for run in p.runs:
        run.font.size = Pt(style["body_size"])


def _render_experience(doc: Document, data: dict[str, Any], style: dict[str, Any]) -> None:
    experience = data.get("experience", [])
    if not experience:
        return
    _add_heading(doc, "Experience", style)
    for exp in experience:
        title = exp.get("title", "")
        company = exp.get("company", "")
        location = exp.get("location", "")
        dates = exp.get("dates", "")
        header = " — ".join(p for p in [title, company] if p)
        if location:
            header = f"{header} ({location})" if header else location
        if dates:
            header = f"{header}  ({dates})" if header else dates
        if header:
            p = doc.add_paragraph()
            run = p.add_run(header)
            run.bold = True
            run.font.size = Pt(style["body_size"])
        for bullet in exp.get("bullets", []):
            _add_bullet(doc, bullet, style)


def _render_education(doc: Document, data: dict[str, Any], style: dict[str, Any]) -> None:
    education = data.get("education", [])
    if not education:
        return
    _add_heading(doc, "Education", style)
    for edu in education:
        line = edu.get("degree", "")
        gpa = edu.get("gpa", "")
        inst = edu.get("institution", "")
        dates = edu.get("dates", "")
        if gpa and line:
            line = f"{line} | CGPA: {gpa}"
        elif gpa:
            line = f"CGPA: {gpa}"
        if inst:
            line = f"{line}, {inst}" if line else inst
        if dates:
            line = f"{line} ({dates})" if line else dates
        if line:
            p = doc.add_paragraph(line)
            for run in p.runs:
                run.font.size = Pt(style["body_size"])


def _render_projects(doc: Document, data: dict[str, Any], style: dict[str, Any]) -> None:
    projects = data.get("projects", [])
    if not projects:
        return
    _add_heading(doc, "Projects", style)
    for proj in projects:
        name = proj.get("name", "")
        if name:
            p = doc.add_paragraph()
            tech = proj.get("tech_stack", "") or proj.get("description", "")
            link = proj.get("link", "")
            title = name
            if tech:
                title = f"{title} | {tech}"
            if link:
                title = f"{title} | {link}"
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(style["body_size"])
        desc = proj.get("description", "")
        if desc and not proj.get("tech_stack"):
            p = doc.add_paragraph(desc)
            for run in p.runs:
                run.font.size = Pt(style["body_size"])
        for bullet in proj.get("bullets", []):
            _add_bullet(doc, bullet, style)


def _render_achievements(doc: Document, data: dict[str, Any], style: dict[str, Any]) -> None:
    achievements = data.get("achievements", [])
    if not achievements:
        return
    _add_heading(doc, "Achievements & Recognition", style)
    for item in achievements:
        _add_bullet(doc, str(item), style)
